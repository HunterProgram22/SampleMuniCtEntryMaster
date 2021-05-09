import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt5.QtWidgets import (
    QApplication, QDialog, QMainWindow, QMessageBox
)
from PyQt5.uic import loadUi
from PyQt5.QtCore import pyqtSlot

from main_window_ui import Ui_MainWindow
from final_judgment_entry_ui import Ui_Dialog
from bond_sheet_ui import Ui_BondSheetDialog

class Window(QMainWindow, Ui_MainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        self.connectSignalsSlots()

    def connectSignalsSlots(self):
        self.actionFinal_Judgment_Entry.triggered.connect(self.FinalJudgmentEntry)
        self.actionBond_Sheet.triggered.connect(self.BondSheet)

    def FinalJudgmentEntry(self):
        dialog = FinalJudgmentEntryDialog(self)
        dialog.exec()

    def BondSheet(self):
        dialog = BondSheetDialog(self)
        dialog.exec()



class FinalJudgmentEntryDialog(QDialog, Ui_Dialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        #loadUi("ui/final_judgment_entry.ui", self)

    def open_doc(self):
        """For templates need to make sure all styles used are part of the template. One solution is the template has all types of
        text and then is deleted when opened."""
        document = Document("entry_template.docx")
        heading = document.add_heading('Final Judgment Entry', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name = document.add_paragraph("The defendant is: {0}".format(self.lineEdit.text()))
        p_case_number = document.add_paragraph("The case number is: {0}".format(self.lineEdit_2.text()))
        p_charge = document.add_paragraph("The charge is: {0}".format(self.comboBox.currentText()))
        p_plea_and_finding = document.add_paragraph("{0} has plead {1} and this court finds {0} {2}.".format(self.lineEdit.text(), self.comboBox_2.currentText(), self.comboBox_3.currentText()))
        p_comm_control_heading = document.add_paragraph()
        p_comm_control_heading.add_run("COMMUNITY CONTROL SANCTIONS").bold = True
        if self.checkBox.isChecked():
            p_comm_control = document.add_paragraph("The defendant is required to be tested for drugs.", style='List Paragraph')
        if self.checkBox_2.isChecked():
            p_comm_control = document.add_paragraph("The defendant's driver's license is revoked.", style='List Paragraph')
        if self.checkBox_3.isChecked():
            p_comm_control = document.add_paragraph("Weekend reporting is required for {0} for the charge of {1}.".format(self.lineEdit.text(), self.comboBox.currentText()), style='List Paragraph')
        if self.checkBox_4.isChecked():
            p_comm_control = document.add_paragraph("GPS Monitoring is required for {0}.".format(self.lineEdit.text()), style='List Paragraph')
        #else:
            #p_comm_control = document.add_paragraph("There are no community control " +\
            #"sanctions for {0}".format(self.lineEdit.text()))
        document.save('demo.docx')


class BondSheetDialog(QDialog, Ui_BondSheetDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setupUi(self)
        #loadUi("ui/final_judgment_entry.ui", self)

    def open_doc(self):
        pass


if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = Window()
    win.show()
    sys.exit(app.exec())
